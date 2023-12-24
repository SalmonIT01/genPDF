import os
from subprocess import Popen
from docx import Document
from docx.shared import Inches,Cm
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert
from docx.shared import Inches, Pt
from io import BytesIO
import fitz  # PyMuPDF

import mammoth



def gen_word():
        
        doc = Document()
        # Add image
        doc.add_picture('crud.jpg', width=Cm(3),height=Cm(3))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # style
        head_style = doc.styles.add_style('HeadStyle',WD_STYLE_TYPE.PARAGRAPH,)
        font = head_style.font
        font.name = 'TH Sarabun New'
        font.size = Pt(18)
        font.bold = True

        BodyStyle = doc.styles.add_style('BodyStyle',WD_STYLE_TYPE.PARAGRAPH,)
        font = BodyStyle.font
        font.name = 'TH Sarabun New'
        font.size = Pt(16)
        font.bold = False


        def boldandunderline(p,x):
            result = p.add_run(x)
            result.bold = True
            result.underline = True

        data = ["ประกาศบัณฑิตวิทยาลัย มหาวิทยาลัยนเรศวร","เรื่อง อนุมัติให้นิสิตระดับปริญญาโทดำเนินการทำวิจัย","ครั้งที่ ๐๐๕/๒๕๖๖\n",]

        for i in data:
            p = doc.add_paragraph(i)
            p.style = head_style
            p.paragraph_format.line_spacing = Inches(0.3)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER




        info = {"name":"นายอนันต์ วรรณศรี",
                "std_code":"๖๓๐๖๒๐๐๕",
                "std_fac":"หลักสูตรปริญญาบริหารธุรกิจมหาบัณฑิต สาขาวิชาการบริหารเทคโนโลยีสารสนเทศเชิงกลยุทธ์",
                # "project_nameTH":"การวางยุทธศาสตร์เมืองอัจฉริยะด้วยระบบออโตเมชั่นด้านการวางผังเมือง และโครงสร้างพื้นฐานขององค์การบริหารส่วนตำบล อำเถอลานกระบือ"
                "name101":"การพัฒนาตลาดการท่องเที่ยวของนักท่องเที่ยวกลุ่มมิลเลนเนียมด้วย แนวคิดความผูกพันทางอารมณ์ผ่านผู้ทรงอิทธิพลในสื่อสังคัง",
                "project_nameEN":"SMART CITY STRATEGIC PLANNING WITH URBAN PLANNING AUTOMATION AND INFRASTRUCTURE OF SUB-DISTRICT ADMINISTRATIVE ORGANIZATIONS IN LAN KRABUE DISTRICT",
                "teacher_name":"ผู้ช่วยศาสตราจารย์ ดร.วศิน เหลี่ยมปรีชา"}

        p = doc.add_paragraph('บัณฑิตวิทยาลัยอนุมัติให้ ')
        p.paragraph_format. first_line_indent = Inches(0.25)
        p.style = BodyStyle
        p.add_run('{0} '.format(info['name'])).bold = True
        p.add_run('รหัสประจำตัว {0} นิสิตระดับปริญญาโท '.format(info['std_code']))
        p.add_run('{0} ดำเนินการทำวิจัยตามโครงร่างวิทยานิพนธ์ที่เสนอ'.format(info['std_fac']))
        # p.add_run(' ดำเนินการทำวิจัยตามโครงร่างวิทยานิพนธ์ที่เสนอ')


        data = ( 
        ('เรื่อง','ภาษาไทย',info['name101']), 
        ('','ภาษาอังกฤษ',info['project_nameEN']), 
        ('','โดยมี',info['teacher_name']+" เป็นประธานที่ปรึกษาวิทยานิพนธ์") 
        ) 

        # Creating a table object 
        table = doc.add_table(rows=0, cols=3) 
        # table.columns[0].width = Cm(1.19)
        # table.columns[1].width = Cm(2.5)
        # table.columns[2].width = Cm(11.19)

        # table.cell(0,0).text = 'เรื่อง'
        # table.cell(0,1).text = 'ภาษาไทย'
        # table.cell(0,2).text = 'SMART CITY STRATEGIC PLANNING WITH URBAN PLANNING AUTOMATION AND INFRASTRUCTURE OF SUB-DISTRICT ADMINISTRATIVE ORGANIZATIONS IN LAN KRABUE DISTRICT'



        for name,title,project in data:
            row = table.add_row().cells 
            
            row[0].text = name
            p = row[1].paragraphs[0]
            p.add_run(title).bold = True
            row[2].text = project

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = BodyStyle # <---- style
                    



        p = doc.add_paragraph("จึงประกาศมาให้ทราบโดยทั่วกัน")
        p.style = BodyStyle


        docx_file = 'demo.docx'
        doc.save(docx_file)
        pdf_file = 'demo.pdf'
        convert(docx_file, pdf_file)
        os.remove(docx_file)
        
        # file_stream = BytesIO()
        # doc.save(file_stream)
        # file_stream.seek(0)
        # result = mammoth.convert_to_html(file_stream)  
        # print(result.value)
        # f = open("demofile3.txt", "w")
        # f.write(result.value)
        # f.close()
        
        pdf_bytesio = None
           
        try:
            pdf_bytesio = pdf_to_bytesio(pdf_file)
        except:
            pass
        
        if pdf_bytesio != None:
            # print(pdf_bytesio.getvalue())
            os.remove(pdf_file)
            return pdf_bytesio.getvalue()

def pdf_to_bytesio(pdf_path):
    pdf_bytesio = BytesIO()
    with open(pdf_path, 'rb') as file:
        pdf_bytesio.write(file.read())
    return pdf_bytesio

